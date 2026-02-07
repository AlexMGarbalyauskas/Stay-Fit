from datetime import date
from docx import Document

OUTPUT_PATH = "Project_Technical_Report.docx"


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_paragraph(doc, text):
    doc.add_paragraph(text)


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def main():
    doc = Document()

    title = "Stay-Fit Project Technical Report"
    doc.add_heading(title, level=0)
    add_paragraph(doc, f"Date: {date.today().isoformat()}")

    add_heading(doc, "Project Idea and Goals", level=1)
    add_paragraph(
        doc,
        "Stay-Fit is a full-stack social fitness web application. The core idea is to combine workout tracking "
        "with social networking and real-time messaging so users can share progress, connect with friends, "
        "and schedule workouts together."
    )

    add_heading(doc, "High-Level Architecture", level=1)
    add_bullets(
        doc,
        [
            "Frontend: React single-page application (SPA).",
            "Backend: Node.js with Express REST API.",
            "Database: SQLite (file-based relational database).",
            "Real-time: Socket.IO for WebSocket communication.",
            "File storage: Local uploads directory served by Express.",
        ],
    )

    add_heading(doc, "Backend Technical Aspects", level=1)
    add_heading(doc, "Framework and Runtime", level=2)
    add_bullets(
        doc,
        [
            "Node.js runtime with Express 5 for HTTP routing and middleware.",
            "Socket.IO for real-time messaging and notification events.",
            "CORS and security headers (Helmet) configured for safe cross-origin access.",
        ],
    )

    add_heading(doc, "Data Layer", level=2)
    add_bullets(
        doc,
        [
            "SQLite database stored as a local file (data.sqlite).",
            "Raw SQL via sqlite3 driver; schema evolved with SQL migration files.",
            "Migration runner applies versioned SQL scripts at startup.",
        ],
    )

    add_heading(doc, "Authentication and Security", level=2)
    add_bullets(
        doc,
        [
            "JWT-based auth for API requests and Socket.IO connections.",
            "Password hashing with bcrypt.",
            "Google OAuth 2.0 via Passport for social login.",
            "Email verification flow using short-lived tokens/codes.",
            "End-to-end message encryption on the client using AES-GCM; encrypted payload stored server-side.",
        ],
    )

    add_heading(doc, "Core Backend Features", level=2)
    add_bullets(
        doc,
        [
            "REST endpoints for users, posts, comments, likes, saves, and notifications.",
            "Friend requests and social graph management.",
            "Direct messaging with optional encrypted content and media attachments.",
            "Workout scheduling with invite and response flow.",
            "Multer-based uploads for images and videos, with type validation.",
        ],
    )

    add_heading(doc, "Frontend Technical Aspects", level=1)
    add_heading(doc, "Framework and Build", level=2)
    add_bullets(
        doc,
        [
            "React 19 with Create React App tooling.",
            "React Router for client-side routing and guarded routes.",
            "Axios for HTTP requests with JWT interceptor.",
            "Socket.IO client for live messaging and notifications.",
        ],
    )

    add_heading(doc, "UI and State Management", level=2)
    add_bullets(
        doc,
        [
            "Tailwind CSS for styling and responsive layout.",
            "Context API for shared state (language, workout reminders, auth state).",
            "Component-based UI for feed, profiles, messaging, and settings.",
        ],
    )

    add_heading(doc, "Key Frontend Features", level=2)
    add_bullets(
        doc,
        [
            "Authentication flows: login, register, OAuth, and email verification.",
            "Social feed with posts, comments, likes, and saved posts.",
            "Real-time chat with emoji reactions and media support.",
            "Calendar and reminders for workout scheduling.",
            "Notifications hub with read/unread controls.",
        ],
    )

    add_heading(doc, "Technology Summary", level=1)
    add_bullets(
        doc,
        [
            "Backend: Node.js, Express, SQLite, Socket.IO, JWT, Passport.",
            "Frontend: React, React Router, Axios, Tailwind CSS, Socket.IO client.",
            "Security: bcrypt hashing, JWT auth, email verification, optional message encryption.",
        ],
    )

    add_heading(doc, "Notes and Assumptions", level=1)
    add_paragraph(
        doc,
        "This report was derived from project technical notes and code structure. README content was minimal, "
        "so the technical notes and source files informed the architecture summary."
    )

    doc.save(OUTPUT_PATH)


if __name__ == "__main__":
    main()
